from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Shading Accent 1"
    for k, v in rows:
        r = table.add_row().cells
        r[0].text = k
        r[1].text = v


def add_findings_table(
    doc: Document,
    findings: list[dict[str, str]],
) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Priority"
    hdr[1].text = "Finding"
    hdr[2].text = "Evidence"
    hdr[3].text = "Impact"
    hdr[4].text = "Suggested fix"
    for f in findings:
        row = table.add_row().cells
        row[0].text = f["priority"]
        row[1].text = f["finding"]
        row[2].text = f["evidence"]
        row[3].text = f["impact"]
        row[4].text = f["fix"]


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    out_path = repo_root / "output" / "doc" / f"QuanCore_Beta_Readiness_Report_{date.today().isoformat()}.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("QuanCore — Beta Launch Ship Readiness Report", level=1)
    doc.add_paragraph(f"Date: {date.today().isoformat()}")

    doc.add_heading("Executive summary", level=2)
    doc.add_paragraph(
        "Overall: READY FOR BETA with conditions. The app now builds cleanly and passes lint/typecheck, "
        "but you must rotate leaked secrets and verify auth + billing flows in the target environment."
    )

    doc.add_heading("Scope & environment", level=2)
    add_kv_table(
        doc,
        [
            ("App location", str(repo_root)),
            ("Framework", "Next.js 16.1.6 (App Router)"),
            ("Runtime", "Node 24.13.0 / npm 11.6.2 (observed)"),
            ("Key services", "Supabase Auth/DB, OpenRouter, Cashfree (sandbox config present)"),
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Evidence-based checks run", level=2)
    doc.add_paragraph("Commands executed and their outcomes:")
    for line in [
        "npm run typecheck  → PASS",
        "npm run lint       → PASS (after fixes)",
        "npm run build      → PASS (after fixes); warning remains about middleware/proxy deprecation",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Top findings (with suggested fixes)", level=2)

    findings: list[dict[str, str]] = [
        {
            "priority": "P0",
            "finding": "Secrets present in repo working tree (.env / .env.local)",
            "evidence": "Supabase anon key, service role key, and OpenRouter key were present in plaintext.",
            "impact": "Credential leakage risk; service-role key can fully compromise data.",
            "fix": "Rotate keys immediately; keep secrets only in deployment environment; commit only .env.example.",
        },
        {
            "priority": "P0",
            "finding": "Production build previously failed on /login (useSearchParams without Suspense)",
            "evidence": "next build error: “useSearchParams() should be wrapped in a suspense boundary at page /login”.",
            "impact": "Cannot produce production build artifact; blocks launch.",
            "fix": "Wrap login page content in <Suspense> (implemented).",
        },
        {
            "priority": "P1",
            "finding": "Lint previously failed due to config rule and unused imports",
            "evidence": "eslint error in next.config.js (require import rule) + unused Button imports.",
            "impact": "CI/quality gate likely fails; increases noise and hides real issues.",
            "fix": "Add eslint override for next.config.js + remove unused imports (implemented).",
        },
        {
            "priority": "P1",
            "finding": "Next.js warns: middleware file convention deprecated; use proxy instead",
            "evidence": "next build warns: “middleware file convention is deprecated… use proxy instead”.",
            "impact": "Forward-compat risk; could break on future upgrades.",
            "fix": "Plan a migration to Next.js proxy convention per Next docs; retest auth gate paths.",
        },
        {
            "priority": "P1",
            "finding": "Supabase Realtime send() fallback warnings observed in dev logs",
            "evidence": "Dev logs show “Realtime send() is automatically falling back to REST API… will be deprecated”.",
            "impact": "Potential future breakage and real-time latency/behavior drift.",
            "fix": "Switch to explicit REST delivery (httpSend) or correct Realtime channel setup; add monitoring.",
        },
        {
            "priority": "P1",
            "finding": "Agent creation may be blocked by plan limit (403 /api/agents) in E2E run report",
            "evidence": "PLAYWRIGHT_RUN_REPORT.md: POST /api/agents → 403, UI: “Plan limit reached…”.",
            "impact": "Beta users may be unable to create expected agents; product feels broken.",
            "fix": "Confirm plan logic for beta cohort; add clearer UI copy + upgrade path; add an admin override for beta.",
        },
        {
            "priority": "P2",
            "finding": "README is generic create-next-app template",
            "evidence": "README.md lacks environment setup, Supabase steps, and beta runbook.",
            "impact": "Onboarding friction for beta testers and teammates.",
            "fix": "Replace with a beta runbook: env vars, Supabase project setup, seed user, billing/webhook setup.",
        },
    ]
    add_findings_table(doc, findings)

    doc.add_paragraph()
    doc.add_heading("What was changed during this review (to unblock beta)", level=2)
    for item in [
        "Fixed production build: wrapped /login page in React Suspense boundary.",
        "Fixed lint: removed unused imports; added eslint override for next.config.js require usage.",
        "Removed plaintext secrets from .env and .env.local; added .env.example template.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Beta launch checklist (today)", level=2)
    checklist = [
        "Rotate Supabase service role key and OpenRouter key; verify no leaked keys remain in shared files.",
        "Set production environment variables in hosting platform (do not rely on local .env files).",
        "Verify login + register for the beta cohort (Supabase email confirmation/rate limit settings).",
        "Verify critical flows: create agent, start swarm run, stop run, view history, settings page.",
        "Verify Cashfree webhook endpoint in the beta environment (if billing is enabled).",
        "Set basic observability: error logging for API routes + a minimal uptime check on / and /login.",
    ]
    for item in checklist:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph("Prepared by: Cursor automated review")

    doc.save(out_path)
    print(str(out_path))


if __name__ == "__main__":
    main()

